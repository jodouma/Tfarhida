from __future__ import annotations

import argparse
import html
import math
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Sequence
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, StyleSheet1, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    HRFlowable,
    Image as RLImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table as RLTable,
    TableStyle,
)

SOURCE_MD = Path("/home/jo/environnements/Tfarhida/docs/report/rapport-pfe-tfarhida-final.md")
DOCX_OUT = Path(
    "/home/jo/environnements/Tfarhida/docs/report/"
    "Rapport_PFE_Tfarhida_Yosra_El_Hadj_Brayek_Wassim_Chommakh.docx"
)
PDF_OUT = Path(
    "/home/jo/environnements/Tfarhida/docs/report/"
    "Rapport_PFE_Tfarhida_Yosra_El_Hadj_Brayek_Wassim_Chommakh.pdf"
)

ACCENT = RGBColor(0x12, 0x3C, 0x73)
TEXT = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)
LIGHT_BORDER = RGBColor(0xD1, 0xD5, 0xDB)
LIGHT_FILL = RGBColor(0xF3, 0xF6, 0xFA)

REPORTLAB_BLUE = colors.HexColor("#123C73")
REPORTLAB_TEXT = colors.HexColor("#1F2937")
REPORTLAB_MUTED = colors.HexColor("#6B7280")
REPORTLAB_FILL = colors.HexColor("#F3F6FA")
REPORTLAB_BORDER = colors.HexColor("#D1D5DB")


@dataclass
class HeadingBlock:
    level: int
    text: str


@dataclass
class ParagraphBlock:
    lines: list[str]


@dataclass
class ListBlock:
    ordered: bool
    items: list[str]


@dataclass
class ImageBlock:
    caption: str
    path: Path
    width_pct: int


@dataclass
class TableBlock:
    headers: list[str]
    rows: list[list[str]]
    caption: str | None


@dataclass
class CodeBlock:
    language: str
    text: str


@dataclass
class QuoteBlock:
    text: str


@dataclass
class PageBreakBlock:
    pass


Block = (
    HeadingBlock
    | ParagraphBlock
    | ListBlock
    | ImageBlock
    | TableBlock
    | CodeBlock
    | QuoteBlock
    | PageBreakBlock
)


IMAGE_RE = re.compile(
    r"!\[(?P<caption>.*?)\]\((?P<path>[^)]+)\)(?:\{\s*width=(?P<width>\d+)%\s*\})?"
)
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")


def normalize_text(text: str) -> str:
    text = text.replace("\u2014", "-")
    text = text.replace("\u2013", "-")
    return text


def strip_markdown_links(text: str) -> str:
    def repl(match: re.Match[str]) -> str:
        label = match.group(1)
        url = match.group(2)
        return url if label == url else f"{label} ({url})"

    return re.sub(r"\[([^\]]+)\]\(([^)]+)\)", repl, text)


def strip_markdown_formatting(text: str) -> str:
    text = strip_markdown_links(text)
    text = text.replace("**", "")
    text = text.replace("`", "")
    return normalize_text(text)


def markdown_inline_to_html(text: str) -> str:
    text = strip_markdown_links(text)
    parts: list[str] = []
    cursor = 0
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    for match in pattern.finditer(text):
        if match.start() > cursor:
            parts.append(html.escape(text[cursor : match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            parts.append(f"<b>{html.escape(token[2:-2])}</b>")
        else:
            parts.append(
                f'<font face="Courier">{html.escape(token[1:-1])}</font>'
            )
        cursor = match.end()
    if cursor < len(text):
        parts.append(html.escape(text[cursor:]))
    return normalize_text("".join(parts))


def iter_markdown_blocks(source: Path) -> list[Block]:
    lines = source.read_text(encoding="utf-8").splitlines()
    try:
        start = next(i for i, line in enumerate(lines) if line.strip() == "# Remerciements")
    except StopIteration as exc:
        raise RuntimeError("Could not find report body start.") from exc

    blocks: list[Block] = []
    i = start
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == r"\newpage":
            blocks.append(PageBreakBlock())
            i += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append(CodeBlock(language=language, text="\n".join(code_lines)))
            i += 1
            continue

        image_match = IMAGE_RE.fullmatch(stripped)
        if image_match:
            rel_path = image_match.group("path")
            width_pct = int(image_match.group("width") or 80)
            blocks.append(
                ImageBlock(
                    caption=strip_markdown_formatting(image_match.group("caption")),
                    path=(source.parent / rel_path).resolve(),
                    width_pct=width_pct,
                )
            )
            i += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            blocks.append(
                HeadingBlock(
                    level=len(heading_match.group(1)),
                    text=strip_markdown_formatting(heading_match.group(2)),
                )
            )
            i += 1
            continue

        if "|" in stripped and i + 1 < len(lines):
            separator = lines[i + 1].strip()
            if separator and set(separator.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")) == set():
                raw_rows = [line]
                i += 2
                while i < len(lines) and lines[i].strip().startswith("|"):
                    raw_rows.append(lines[i])
                    i += 1
                caption = None
                probe = i
                while probe < len(lines) and not lines[probe].strip():
                    probe += 1
                if probe < len(lines) and lines[probe].strip().startswith(":"):
                    caption = strip_markdown_formatting(lines[probe].strip()[1:].strip())
                    i = probe + 1
                headers = split_table_row(raw_rows[0])
                rows = [split_table_row(row) for row in raw_rows[1:]]
                blocks.append(TableBlock(headers=headers, rows=rows, caption=caption))
                continue

        if stripped.startswith("> "):
            quote_lines = [stripped[2:].strip()]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("> "):
                quote_lines.append(lines[i].strip()[2:].strip())
                i += 1
            blocks.append(QuoteBlock(text=strip_markdown_formatting(" ".join(quote_lines))))
            continue

        if stripped.startswith("- "):
            items: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(strip_markdown_formatting(lines[i].strip()[2:].strip()))
                i += 1
            blocks.append(ListBlock(ordered=False, items=items))
            continue

        ordered_match = ORDERED_RE.match(stripped)
        if ordered_match:
            items = []
            while i < len(lines):
                current = ORDERED_RE.match(lines[i].strip())
                if not current:
                    break
                items.append(strip_markdown_formatting(current.group(1)))
                i += 1
            blocks.append(ListBlock(ordered=True, items=items))
            continue

        para_lines: list[str] = []
        while i < len(lines):
            current = lines[i]
            current_stripped = current.strip()
            if not current_stripped:
                break
            if current_stripped == r"\newpage":
                break
            if current_stripped.startswith("```"):
                break
            if IMAGE_RE.fullmatch(current_stripped):
                break
            if HEADING_RE.match(current_stripped):
                break
            if current_stripped.startswith("> "):
                break
            if current_stripped.startswith("- "):
                break
            if ORDERED_RE.match(current_stripped):
                break
            if "|" in current_stripped and i + 1 < len(lines):
                sep = lines[i + 1].strip()
                if sep and set(sep.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")) == set():
                    break
            para_lines.append(current.rstrip())
            i += 1
        blocks.append(ParagraphBlock(lines=para_lines))
    return blocks


def split_table_row(row: str) -> list[str]:
    cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
    return [strip_markdown_formatting(cell) for cell in cells]


def configure_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(22)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 14, ACCENT),
        ("Heading 3", 12, TEXT),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(6)

    if "Caption" not in doc.styles:
        caption = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["Caption"]
    caption.font.name = "Arial"
    caption.font.size = Pt(9.5)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if "Quote" not in doc.styles:
        quote = doc.styles.add_style("Quote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote = doc.styles["Quote"]
    quote.font.name = "Arial"
    quote.font.size = Pt(10.5)
    quote.font.italic = True
    quote.font.color.rgb = MUTED
    quote._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    quote._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    quote.paragraph_format.left_indent = Inches(0.35)
    quote.paragraph_format.space_after = Pt(8)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        text = OxmlElement("w:t")
        text.text = "1"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(separate)
        run._r.append(text)
        run._r.append(end)


def add_cover(doc: Document) -> None:
    logo_path = SOURCE_MD.parent.parent / "assets/report/logos/leaders-university-logo.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run()
    run.add_picture(str(logo_path), width=Inches(4.7))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Rapport de Projet de Fin d'Etudes")
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("Tfarhida - Application web de mini-jeux sociaux tunisiens")
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = TEXT

    grid = doc.add_table(rows=4, cols=2)
    grid.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid.autofit = False
    pairs = [
        ("Etudiantes", "Yosra El Hadj Brayek\nWassim Chommakh"),
        ("Encadrante", "Madame Imen Herzi"),
        ("Etablissement", "Leaders University - Nabeul"),
        ("Annee universitaire", "2025-2026"),
    ]
    widths = [Inches(1.9), Inches(4.2)]
    for row, (label, value) in zip(grid.rows, pairs):
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
        row.cells[0].text = label
        row.cells[1].text = value
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(4)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11.5)
                run.font.color.rgb = TEXT
                if idx == 0:
                    run.bold = True
            if idx == 0:
                set_cell_shading(cell, "F3F6FA")

    doc.add_paragraph()
    doc.add_page_break()


def add_docx_inline_runs(paragraph, text: str) -> None:
    working = strip_markdown_links(text)
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    cursor = 0
    for match in pattern.finditer(working):
        if match.start() > cursor:
            run = paragraph.add_run(normalize_text(working[cursor : match.start()]))
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.font.color.rgb = TEXT
        token = match.group(0)
        content = token[2:-2] if token.startswith("**") else token[1:-1]
        run = paragraph.add_run(normalize_text(content))
        if token.startswith("**"):
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(11)
        else:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        run.font.color.rgb = TEXT
        cursor = match.end()
    if cursor < len(working):
        run = paragraph.add_run(normalize_text(working[cursor:]))
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT


def join_paragraph_lines(lines: Sequence[str]) -> str:
    fragments: list[str] = []
    for index, line in enumerate(lines):
        stripped = line.rstrip()
        cleaned = normalize_text(stripped.rstrip())
        if index == 0:
            fragments.append(cleaned)
            continue
        if lines[index - 1].endswith("  "):
            fragments.append("\n" + cleaned)
        else:
            fragments.append(" " + cleaned.lstrip())
    return "".join(fragments)


def get_doc_content_width_inches(doc: Document) -> float:
    section = doc.sections[-1]
    return (section.page_width - section.left_margin - section.right_margin) / 914400.0


def render_docx(blocks: Sequence[Block], output_path: Path) -> None:
    doc = Document()
    configure_doc_styles(doc)
    add_cover(doc)

    last_heading = ""
    for block in blocks:
        if isinstance(block, HeadingBlock):
            paragraph = doc.add_paragraph(style=f"Heading {min(block.level, 3)}")
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(14 if block.level == 1 else 10)
            add_docx_inline_runs(paragraph, block.text)
            last_heading = block.text
            continue

        if isinstance(block, ParagraphBlock):
            text = join_paragraph_lines(block.lines)
            if not text.strip():
                continue
            paragraph = doc.add_paragraph(style="Normal")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if "\n" in text:
                for idx, part in enumerate(text.split("\n")):
                    if idx:
                        paragraph.add_run().add_break(WD_BREAK.LINE)
                    add_docx_inline_runs(paragraph, part)
            else:
                add_docx_inline_runs(paragraph, text)
            continue

        if isinstance(block, QuoteBlock):
            if last_heading == "Table des matieres":
                paragraph = doc.add_paragraph(block.text, style="Quote")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue
            paragraph = doc.add_paragraph(block.text, style="Quote")
            continue

        if isinstance(block, ListBlock):
            style = "List Number" if block.ordered else "List Bullet"
            for item in block.items:
                paragraph = doc.add_paragraph(style=style)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_docx_inline_runs(paragraph, item)
            continue

        if isinstance(block, ImageBlock):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            width_inches = get_doc_content_width_inches(doc) * (block.width_pct / 100.0)
            paragraph.add_run().add_picture(str(block.path), width=Inches(width_inches))
            if block.caption:
                caption = doc.add_paragraph(block.caption, style="Caption")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if isinstance(block, TableBlock):
            table = doc.add_table(rows=1, cols=len(block.headers))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            header_cells = table.rows[0].cells
            for idx, header in enumerate(block.headers):
                header_cells[idx].text = header
                set_cell_shading(header_cells[idx], "EAF0F8")
                header_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in header_cells[idx].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.name = "Arial"
                        run.font.size = Pt(10.5)
            for row_values in block.rows:
                row_cells = table.add_row().cells
                for idx, value in enumerate(row_values):
                    row_cells[idx].text = value
                    row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph in row_cells[idx].paragraphs:
                        paragraph.paragraph_format.space_after = Pt(3)
                        paragraph.alignment = (
                            WD_ALIGN_PARAGRAPH.CENTER
                            if len(value) < 28
                            else WD_ALIGN_PARAGRAPH.LEFT
                        )
                        for run in paragraph.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(10)
                            run.font.color.rgb = TEXT
            if block.caption:
                doc.add_paragraph(block.caption, style="Caption")
            continue

        if isinstance(block, CodeBlock):
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.rows[0].cells[0]
            set_cell_shading(cell, "F8FAFC")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            for idx, line in enumerate(block.text.splitlines()):
                if idx:
                    paragraph.add_run().add_break(WD_BREAK.LINE)
                run = paragraph.add_run(line)
                run.font.name = "Courier New"
                run.font.size = Pt(8.8)
                run.font.color.rgb = TEXT
            doc.add_paragraph()
            continue

        if isinstance(block, PageBreakBlock):
            doc.add_page_break()

    add_page_number_footer(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def build_pdf_styles() -> StyleSheet1:
    base = getSampleStyleSheet()
    styles = StyleSheet1()

    styles.add(
        ParagraphStyle(
            name="Body",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=15,
            textColor=REPORTLAB_TEXT,
            spaceAfter=8,
            alignment=TA_JUSTIFY,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TitleCover",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=28,
            textColor=REPORTLAB_BLUE,
            alignment=TA_CENTER,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SubtitleCover",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=22,
            textColor=REPORTLAB_TEXT,
            alignment=TA_CENTER,
            spaceAfter=18,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Heading1Report",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15.5,
            leading=21,
            textColor=REPORTLAB_BLUE,
            spaceBefore=8,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Heading2Report",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13.2,
            leading=18,
            textColor=REPORTLAB_BLUE,
            spaceBefore=6,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Heading3Report",
            parent=base["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=11.6,
            leading=15,
            textColor=REPORTLAB_TEXT,
            spaceBefore=4,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Caption",
            parent=base["Italic"],
            fontName="Helvetica-Oblique",
            fontSize=8.8,
            leading=11,
            textColor=REPORTLAB_MUTED,
            alignment=TA_CENTER,
            spaceAfter=8,
            spaceBefore=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Quote",
            parent=base["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=10.2,
            leading=14,
            leftIndent=18,
            textColor=REPORTLAB_MUTED,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SmallCenter",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10.2,
            leading=13,
            alignment=TA_CENTER,
            textColor=REPORTLAB_TEXT,
            spaceAfter=4,
        )
    )
    return styles


def get_image_size(path: Path, width_points: float) -> tuple[float, float]:
    with PILImage.open(path) as img:
        width, height = img.size
    ratio = height / width
    return width_points, width_points * ratio


def render_pdf(blocks: Sequence[Block], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=2.0 * cm,
        rightMargin=2.0 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.5 * cm,
        title="Rapport PFE Tfarhida",
        author="Yosra El Hadj Brayek; Wassim Chommakh",
    )
    styles = build_pdf_styles()
    story: list = []

    logo_path = SOURCE_MD.parent.parent / "assets/report/logos/leaders-university-logo.png"
    logo_width = 11.5 * cm
    logo_w, logo_h = get_image_size(logo_path, logo_width)
    story.extend(
        [
            Spacer(1, 1.8 * cm),
            RLImage(str(logo_path), width=logo_w, height=logo_h, hAlign="CENTER"),
            Spacer(1, 1.0 * cm),
            Paragraph("Rapport de Projet de Fin d'Etudes", styles["TitleCover"]),
            Paragraph(
                "Tfarhida - Application web de mini-jeux sociaux tunisiens",
                styles["SubtitleCover"],
            ),
        ]
    )

    metadata = [
        ("Etudiantes", "Yosra El Hadj Brayek<br/>Wassim Chommakh"),
        ("Encadrante", "Madame Imen Herzi"),
        ("Etablissement", "Leaders University - Nabeul"),
        ("Annee universitaire", "2025-2026"),
    ]
    meta_data = [
        [
            Paragraph(f"<b>{html.escape(label)}</b>", styles["Body"]),
            Paragraph(value, styles["Body"]),
        ]
        for label, value in metadata
    ]
    meta_table = RLTable(meta_data, colWidths=[4.2 * cm, 10.6 * cm], hAlign="CENTER")
    meta_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), REPORTLAB_FILL),
                ("BOX", (0, 0), (-1, -1), 0.5, REPORTLAB_BORDER),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, REPORTLAB_BORDER),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend([meta_table, Spacer(1, 1.0 * cm), HRFlowable(width="60%", color=REPORTLAB_BORDER), PageBreak()])

    last_heading = ""
    for block in blocks:
        if isinstance(block, HeadingBlock):
            style_name = {
                1: "Heading1Report",
                2: "Heading2Report",
                3: "Heading3Report",
            }.get(block.level, "Heading3Report")
            story.append(Paragraph(markdown_inline_to_html(block.text), styles[style_name]))
            last_heading = block.text
            continue

        if isinstance(block, ParagraphBlock):
            text = join_paragraph_lines(block.lines).replace("\n", "<br/>")
            if text.strip():
                story.append(Paragraph(markdown_inline_to_html(text), styles["Body"]))
            continue

        if isinstance(block, QuoteBlock):
            story.append(Paragraph(markdown_inline_to_html(block.text), styles["Quote"]))
            continue

        if isinstance(block, ListBlock):
            flowables = [
                ListItem(Paragraph(markdown_inline_to_html(item), styles["Body"]))
                for item in block.items
            ]
            story.append(
                ListFlowable(
                    flowables,
                    bulletType="1" if block.ordered else "bullet",
                    leftIndent=18,
                )
            )
            story.append(Spacer(1, 0.12 * cm))
            continue

        if isinstance(block, ImageBlock):
            usable_width = doc.width * (block.width_pct / 100.0)
            image_w, image_h = get_image_size(block.path, usable_width)
            story.append(
                KeepTogether(
                    [
                        RLImage(str(block.path), width=image_w, height=image_h, hAlign="CENTER"),
                        Paragraph(html.escape(block.caption), styles["Caption"]),
                    ]
                )
            )
            continue

        if isinstance(block, TableBlock):
            data = [
                [Paragraph(markdown_inline_to_html(cell), styles["Body"]) for cell in block.headers]
            ]
            for row in block.rows:
                data.append(
                    [Paragraph(markdown_inline_to_html(cell), styles["Body"]) for cell in row]
                )
            total_cols = len(block.headers)
            col_width = doc.width / max(total_cols, 1)
            table = RLTable(
                data,
                colWidths=[col_width] * total_cols,
                repeatRows=1,
                hAlign="CENTER",
            )
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF0F8")),
                        ("GRID", (0, 0), (-1, -1), 0.45, REPORTLAB_BORDER),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(table)
            if block.caption:
                story.append(Paragraph(html.escape(block.caption), styles["Caption"]))
            else:
                story.append(Spacer(1, 0.15 * cm))
            continue

        if isinstance(block, CodeBlock):
            story.append(
                Preformatted(
                    normalize_text(block.text),
                    ParagraphStyle(
                        "Code",
                        fontName="Courier",
                        fontSize=8.3,
                        leading=10,
                        textColor=REPORTLAB_TEXT,
                        backColor=REPORTLAB_FILL,
                        borderColor=REPORTLAB_BORDER,
                        borderWidth=0.4,
                        borderPadding=6,
                        leftIndent=8,
                        rightIndent=8,
                        spaceBefore=4,
                        spaceAfter=8,
                    ),
                )
            )
            continue

        if isinstance(block, PageBreakBlock):
            story.append(PageBreak())

    def draw_page_number(canvas, pdf_doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(REPORTLAB_MUTED)
        canvas.drawCentredString(A4[0] / 2, 1.0 * cm, f"{canvas.getPageNumber()}")
        canvas.restoreState()

    doc.build(story, onFirstPage=draw_page_number, onLaterPages=draw_page_number)


def set_docx_metadata(docx_path: Path) -> None:
    tmp_dir = docx_path.parent / ".tmp_report_docx_meta"
    if tmp_dir.exists():
        for child in tmp_dir.rglob("*"):
            if child.is_file():
                child.unlink()
        for child in sorted(tmp_dir.glob("**/*"), reverse=True):
            if child.is_dir():
                child.rmdir()
    tmp_dir.mkdir(parents=True, exist_ok=True)
    try:
        with zipfile.ZipFile(docx_path, "r") as archive:
            archive.extractall(tmp_dir)
        core = tmp_dir / "docProps" / "core.xml"
        ns = {
            "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
            "dc": "http://purl.org/dc/elements/1.1/",
        }
        tree = ET.parse(core)
        root = tree.getroot()
        creator = root.find("dc:creator", ns)
        if creator is not None:
            creator.text = "Yosra El Hadj Brayek; Wassim Chommakh"
        title = root.find("dc:title", ns)
        if title is not None:
            title.text = "Rapport PFE Tfarhida"
        tree.write(core, encoding="utf-8", xml_declaration=True)

        with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for file_path in tmp_dir.rglob("*"):
                if file_path.is_file():
                    archive.write(file_path, file_path.relative_to(tmp_dir).as_posix())
    finally:
        if tmp_dir.exists():
            for file_path in sorted(tmp_dir.rglob("*"), reverse=True):
                if file_path.is_file():
                    file_path.unlink()
                else:
                    file_path.rmdir()
            tmp_dir.rmdir()


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate final report DOCX and PDF artifacts.")
    parser.add_argument("--source", type=Path, default=SOURCE_MD)
    parser.add_argument("--docx", type=Path, default=DOCX_OUT)
    parser.add_argument("--pdf", type=Path, default=PDF_OUT)
    args = parser.parse_args()

    blocks = iter_markdown_blocks(args.source)
    render_docx(blocks, args.docx)
    set_docx_metadata(args.docx)
    render_pdf(blocks, args.pdf)
    print(f"[OK] DOCX generated: {args.docx}")
    print(f"[OK] PDF generated: {args.pdf}")


if __name__ == "__main__":
    main()
